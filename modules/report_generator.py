def generate_deterministic_summary(
    organization_name: str,
    overall_score: float,
    maturity: str,
    gaps: list[tuple[str, float]],
    top_use_cases: list[dict],
    critical_risk_count: int,
) -> dict:
    gap_names = [category for category, _ in gaps]
    use_case_names = [
        use_case["name"]
        for use_case in top_use_cases[:3]
    ]

    current_state = (
        f"{organization_name} received an overall AI readiness score of "
        f"{overall_score:.2f} out of 5, corresponding to the "
        f"{maturity} maturity level."
    )

    risk_statement = (
        f"The initial assessment identified {critical_risk_count} "
        "critical risks requiring executive review."
        if critical_risk_count
        else "No critical risks were identified in the initial assessment."
    )

    return {
        "current_state": current_state,
        "key_gaps": gap_names,
        "priority_use_cases": use_case_names,
        "risk_summary": risk_statement,
        "immediate_actions": [
            "Assign accountable owners for the selected use case.",
            "Confirm success metrics and baseline performance.",
            "Resolve critical data, governance, and risk prerequisites.",
        ],
    }
def prepare_llm_payload(
    org_profile: dict,
    readiness_results: dict,
    use_cases: list[dict],
    risk_register: list[dict],
    roadmap_actions: list[dict]
) -> dict:
    """
    Sanitizes and aggregates structured data for LLM narrative generation.
    Strictly excludes PII, individual names, and confidential free text.
    """
    # 1. Industry & Size Category Only (Excludes explicit company/person names)
    sanitized_org = {
        "industry": org_profile.get("industry", "Unspecified"),
        "company_size": org_profile.get("company_size", "Unspecified"),
        "ai_maturity_target": org_profile.get("ai_maturity_target", "Unspecified")
    }

    # 2. Aggregated Readiness Metrics
    sanitized_readiness = {
        "overall_score": readiness_results.get("overall_score", 0.0),
        "maturity_level": readiness_results.get("maturity_level", "Unassessed"),
        "category_scores": readiness_results.get("category_scores", {}),
        "top_gaps": [cat for cat, _ in readiness_results.get("top_gaps", [])]
    }

    # 3. High-level Use Case Metadata (Excludes custom free-text assumptions)
    sorted_ucs = sorted(use_cases, key=lambda x: x.get("priority_score", 0), reverse=True)[:3]
    sanitized_use_cases = [
        {
            "name": uc.get("name"),
            "solution_type": uc.get("solution_type"),
            "classification": uc.get("classification"),
            "priority_score": uc.get("priority_score")
        }
        for uc in sorted_ucs
    ]

    # 4. Aggregate Risk Profile (Counts only, no individual risk owners or PII)
    critical_count = sum(1 for r in risk_register if r.get("severity") == "Critical")
    high_count = sum(1 for r in risk_register if r.get("severity") == "High")
    med_count = sum(1 for r in risk_register if r.get("severity") == "Medium")
    low_count = sum(1 for r in risk_register if r.get("severity") == "Low")

    sanitized_risks = {
        "total_risks": len(risk_register),
        "critical_count": critical_count,
        "high_count": high_count,
        "medium_count": med_count,
        "low_count": low_count
    }

    # 5. High-level Roadmap Action Items (Excludes assigned owner names)
    sanitized_roadmap = [
        {
            "phase": act.get("phase"),
            "trigger": act.get("trigger"),
            "action": act.get("action")
        }
        for act in roadmap_actions
    ]

    return {
        "organization": sanitized_org,
        "readiness": sanitized_readiness,
        "priority_use_cases": sanitized_use_cases,
        "risk_summary": sanitized_risks,
        "key_roadmap_adjustments": sanitized_roadmap
    }
import os
from openai import OpenAI
from modules.models import ExecutiveSummary


def generate_ai_summary(payload: dict) -> ExecutiveSummary:
    """
    Calls OpenAI using structured outputs to generate an executive summary.
    Requires OPENAI_API_KEY set in environment variables or Streamlit secrets.
    """
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        raise ValueError("OPENAI_API_KEY is not set.")

    client = OpenAI(api_key=api_key)

    prompt = f"""
    You are an expert enterprise AI strategist. Analyze the following sanitized organizational assessment data 
    and synthesize an executive-level summary narrative.

    Sanitized Assessment Data:
    {payload}

    Be concise, practical, and objective.
    """

    completion = client.beta.chat.completions.parse(
        model="gpt-4o-mini",
        messages=[
            {"role": "system", "content": "Synthesize enterprise AI assessment results into structured executive insights."},
            {"role": "user", "content": prompt},
        ],
        response_format=ExecutiveSummary,
    )

    return completion.choices[0].message.parsed
from io import BytesIO
from docx import Document


def create_executive_report(report_data: dict) -> BytesIO:
    document = Document()

    # 1. Cover Page / Header
    document.add_heading("1AUM AI Deployment Readiness Report", level=0)
    document.add_paragraph(f"Organization: {report_data.get('organization_name', 'Organization')}")
    document.add_paragraph(f"Generated Date: {report_data.get('date', 'N/A')}")
    document.add_page_break()

    # 2. Executive Summary
    document.add_heading("Executive Summary", level=1)
    exec_sum = report_data.get("executive_summary", {})
    document.add_paragraph(exec_sum.get("current_state", ""))

    document.add_heading("Key Readiness Gaps", level=2)
    for gap in exec_sum.get("key_gaps", []):
        document.add_paragraph(gap, style="List Bullet")

    # 3. Organization Profile
    document.add_heading("Organization Profile", level=1)
    org_p = report_data.get("org_profile", {})
    document.add_paragraph(f"Industry: {org_p.get('industry', 'N/A')}")
    document.add_paragraph(f"Company Size: {org_p.get('company_size', 'N/A')}")
    document.add_paragraph(f"Maturity Target: {org_p.get('ai_maturity_target', 'N/A')}")

    # 4. AI Readiness Results
    document.add_heading("AI Readiness Results", level=1)
    readiness = report_data.get("readiness_results", {})
    document.add_paragraph(f"Overall Score: {readiness.get('overall_score', 0.0):.2f} / 5.0")
    document.add_paragraph(f"Maturity Level: {readiness.get('maturity_level', 'Unassessed')}")

    # 5. Category Analysis
    document.add_heading("Category Analysis", level=1)
    cat_scores = readiness.get("category_scores", {})
    for cat, score in cat_scores.items():
        document.add_paragraph(f"{cat}: {score:.2f} / 5.0")

    # 6. Priority Use Cases
    document.add_heading("Priority Use Cases", level=1)
    for use_case in report_data.get("top_use_cases", []):
        document.add_heading(use_case.get("name", "Unnamed Use Case"), level=2)
        document.add_paragraph(f"Priority Score: {use_case.get('priority_score', 'N/A')}")
        document.add_paragraph(f"Classification: {use_case.get('classification', 'N/A')}")
        document.add_paragraph(f"Solution Type: {use_case.get('solution_type', 'N/A')}")
        document.add_paragraph(use_case.get("recommendation_reason", ""))

    # 7. AI Risk Register
    document.add_heading("AI Risk Register Summary", level=1)
    risks = report_data.get("risk_register", [])
    if risks:
        for r in risks:
            document.add_paragraph(
                f"[{r.get('risk_id')}] {r.get('category')} - Severity: {r.get('severity')} (Score: {r.get('risk_score')})"
            )
            document.add_paragraph(f"Mitigation: {r.get('mitigation')}", style="List Bullet")
    else:
        document.add_paragraph("No active risks logged.")

    # 8. 90-Day Roadmap
    document.add_heading("90-Day Implementation Roadmap", level=1)
    roadmap_actions = report_data.get("roadmap_actions", [])
    if roadmap_actions:
        for act in roadmap_actions:
            document.add_paragraph(
                f"[{act.get('phase')}] {act.get('action')} (Trigger: {act.get('trigger')})",
                style="List Bullet"
            )
    else:
        document.add_paragraph("Standard 90-day milestone track applies.")

    # 9. Recommended Next Steps
    document.add_heading("Recommended Next Steps", level=1)
    for act in exec_sum.get("immediate_actions", []):
        document.add_paragraph(act, style="List Bullet")

    # 10. Methodology
    document.add_heading("Methodology", level=1)
    document.add_paragraph(
        "Readiness, prioritization, maturity, and risk scores are calculated through "
        "deterministic scoring rules based on user-provided organizational inputs."
    )

    # 11. Limitations and Disclaimer
    document.add_heading("Limitations and Disclaimer", level=1)
    document.add_paragraph(
        "Results provided in this report are directional and generated for decision-support purposes. "
        "All recommendations and risk controls require review and formal validation by accountable enterprise leads."
    )

    # Save strictly in memory buffer
    file_buffer = BytesIO()
    document.save(file_buffer)
    file_buffer.seek(0)

    return file_buffer